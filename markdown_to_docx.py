# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from markdown_it import MarkdownIt

def add_formatted_text(p, inline_token):
    """
    将一个包含格式的 inline token 添加到 Word 段落中。
    支持加粗和行内代码。
    """
    # 如果token没有子元素，直接添加内容
    if not inline_token.children:
        p.add_run(inline_token.content)
        return

    is_bold = False
    for child in inline_token.children:
        if child.type == 'strong_open':
            is_bold = True
        elif child.type == 'strong_close':
            is_bold = False
        elif child.type == 'code_inline':
            run = p.add_run(child.content)
            run.font.name = 'Courier New'
            # 同时为东亚字符设置字体以确保兼容性
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')
        elif child.type == 'text':
            run = p.add_run(child.content)
            if is_bold:
                run.bold = True

def setup_document_styles(document):
    """设置文档的基本样式"""
    # 设置默认中文字体
    document.styles['Normal'].font.name = '宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 定义代码块样式
    style = document.styles.add_style('CodeStyle', 1)
    style.font.name = 'Courier New'
    style.font.size = Pt(10)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')
    
    # 设置默认表格样式的字体
    table_style = document.styles['Table Grid']
    table_style.font.name = '宋体'
    table_style.font.size = Pt(10)
    table_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def markdown_to_docx(md_text, output_filename):
    """将Markdown字符串转换为.docx文件"""
    document = Document()
    setup_document_styles(document)
    md = MarkdownIt()
    
    tokens = md.parse(md_text)
    
    # 用于解析表格的状态变量
    in_table = False
    in_thead = False
    in_tbody = False
    table = None
    current_row_cells = None
    cell_idx = 0

    for i, token in enumerate(tokens):
        # 标题
        if token.type == 'heading_open':
            level = int(token.tag[1])
            content = tokens[i+1].content.strip()
            document.add_heading(content, level=level)
        
        # 段落
        elif token.type == 'paragraph_open':
            if not in_table:
                p = document.add_paragraph()
                add_formatted_text(p, tokens[i+1])

        # 代码块
        elif token.type == 'fence':
            p = document.add_paragraph(style='CodeStyle')
            p.add_run(token.content.strip())
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)

        # 列表
        elif token.type == 'list_item_open':
            p = document.add_paragraph(style='List Bullet')
            add_formatted_text(p, tokens[i+1])

        # --- 表格状态机逻辑 ---
        elif token.type == 'table_open':
            in_table = True
        
        elif token.type == 'thead_open':
            in_thead = True
        
        elif token.type == 'tbody_open':
            in_tbody = True
        
        elif token.type == 'tr_open':
            cell_idx = 0
            if in_thead:
                # 通过扫描表头行来确定列数并创建表格
                cols = 0
                j = i + 1
                while j < len(tokens) and tokens[j].type != 'tr_close':
                    if tokens[j].type == 'th_open':
                        cols += 1
                    j += 1
                if cols > 0:
                    table = document.add_table(rows=1, cols=cols)
                    table.style = 'Table Grid'
                    current_row_cells = table.rows[0].cells
            elif in_tbody:
                if table is not None:
                    current_row_cells = table.add_row().cells

        elif token.type == 'th_open' or token.type == 'td_open':
            if current_row_cells and cell_idx < len(current_row_cells):
                cell = current_row_cells[cell_idx]
                cell.text = '' # 清空单元格默认段落
                p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                add_formatted_text(p, tokens[i+1])
            cell_idx += 1

        elif token.type == 'tr_close':
            current_row_cells = None
        
        elif token.type == 'thead_close':
            in_thead = False
        
        elif token.type == 'tbody_close':
            in_tbody = False
            
        elif token.type == 'table_close':
            in_table = False
            table = None
            document.add_paragraph() # 在表格后添加一个空行以增加间距

    try:
        document.save(output_filename)
        print(f"文档 '{output_filename}' 已成功生成！")
    except Exception as e:
        print(f"保存文件时出错: {e}")

# --- 主程序入口 ---
if __name__ == "__main__":
    # 使用您提供的最新的Markdown文本
    markdown_content = """
# 信息检索 API 文档

**版本:** 1.1.0
**基础URL:** `http://127.0.0.1:8000`

## 简介

本API服务提供了一个完整的文档信息检索解决方案，包括文档上传、向量化存储和智能检索功能。用户可以通过上传结构化或非结构化的文本文件来构建知识库，并通过自然语言查询来检索最相关的信息。

后端利用Elasticsearch进行高效的混合检索（关键词+向量），并结合先进的重排模型（Reranker）对召回结果进行二次排序，以确保返回结果的精准性。

## 1. 智能检索

### `POST /search`

此端点用于执行文档的智能检索。它接收一个自然语言查询，并通过“召回+精排”的两阶段流程返回最相关的文档列表。

#### 请求体 (Request Body)

**Content-Type:** `application/json`

| 字段 | 类型 | 是否必需 | 描述 |
| :--- | :--- | :--- | :--- |
| `query` | string | 是 | 用户的自然语言查询文本。 |
| `top_k` | integer | 否 | 希望返回的重排后结果数量。默认值为 `3`，取值范围 `1-20`。 |

**示例:**
```json
{
  "query": "什么是向量数据库？",
  "top_k": 5
}
```

#### 响应 (Response)

**200 OK - 成功**

| 字段 | 类型 | 描述 |
| :--- | :--- | :--- |
| `results` | array | 一个包含检索结果对象的列表，按相关性分数从高到低排序。 |
| `results[].score` | float | 文档与查询的相关性分数。 |
| `results[].text` | string | 匹配到的文档内容。 |
| `results[].metadata` | object | 文档的元数据。 |

**示例:**
```json
{
  "results": [
    {
      "score": 0.987,
      "text": "第二行介绍了什么是向量数据库。",
      "metadata": {
        "source": "sample_data.txt",
        "line_number": 2
      }
    },
    {
      "score": 0.85,
      "text": "Elasticsearch不仅是搜索引擎，还能存储和分析向量数据。",
      "metadata": {
        "source": "tech_blog_post.html",
        "mermaid_txt": ""
      }
    }
  ]
}
```

**错误响应**
* `503 Service Unavailable`: 如果后端的核心模型（如Embedding或Reranker模型）未能成功加载。
* `500 Internal Server Error`: 如果在检索过程中发生其他内部错误。

#### cURL 示例
```bash
curl -X POST "[http://127.0.0.1:8000/search](http://127.0.0.1:8000/search)" \
-H "Content-Type: application/json" \
-d '{
  "query": "介绍一下图表工具",
  "top_k": 2
}'
```

## 2. 数据管理

### `POST /upload_txt`

上传一个纯文本文件 (`.txt`)。系统会按行读取文件内容，将每一行作为一个独立的文档进行向量化并存入知识库。

#### 请求体 (Request Body)

**Content-Type:** `multipart/form-data`

| 参数 | 类型 | 是否必需 | 描述 |
| :--- | :--- | :--- | :--- |
| `file` | file | 是 | 要上传的 `.txt` 文件。 |

#### 响应 (Response)

**200 OK - 成功**

| 字段 | 类型 | 描述 |
| :--- | :--- | :--- |
| `message` | string | 操作成功的提示信息。 |
| `filename` | string | 已上传文件的名称。 |
| `items_processed` | integer | 文件中被成功处理并入库的非空行数。 |

**示例:**
```json
{
  "message": "文件 'sample_data.txt' 上传并处理成功。",
  "filename": "sample_data.txt",
  "items_processed": 3
}
```

**错误响应**
* `400 Bad Request`: 如果上传的不是 `.txt` 文件，或者文件中不包含任何有效内容。
* `503 Service Unavailable`: 如果后端核心模型未加载。
* `500 Internal Server Error`: 文件处理过程中发生其他错误。

#### cURL 示例
```bash
curl -X POST "[http://127.0.0.1:8000/upload_txt](http://127.0.0.1:8000/upload_txt)" \
-H "Content-Type: multipart/form-data" \
-F "file=@/path/to/your/document.txt"
```

### `POST /upload_json`

上传一个结构化的 JSON 文件 (`.json`)。该文件必须包含一个对象列表，每个对象代表一个待入库的文档，包含文本内容和自定义元数据。

#### 请求体 (Request Body)

**Content-Type:** `multipart/form-data`

| 参数 | 类型 | 是否必需 | 描述 |
| :--- | :--- | :--- | :--- |
| `file` | file | 是 | 要上传的 `.json` 文件。 |

#### JSON 文件内容格式

上传的JSON文件内容必须是一个列表 (`array`)，列表中的每个元素都是一个具有以下结构的JSON对象 (`object`)：

| 字段 | 类型 | 是否必需 | 描述 |
| :--- | :--- | :--- | :--- |
| `text` | string | 是 | 文档的原始内容。 |
| `metadata` | object | 是 | 与文档相关的元数据。 |
| `metadata.source` | string | 是 | (推荐) 文档的来源标识。 |
| `metadata.mermaid_txt` | string | 是 | (推荐) 与文档相关的Mermaid图表定义文本。 |

**JSON 文件示例 (`my_data.json`):**
```json
[
  {
    "text": "Mermaid是一种基于文本的图表绘制工具。",
    "metadata": {
      "source": "manual_doc_1.md",
      "mermaid_txt": "graph TD; A-->B;"
    }
  },
  {
    "text": "Elasticsearch不仅是搜索引擎，还能存储和分析向量数据。",
    "metadata": {
      "source": "tech_blog_post.html",
      "mermaid_txt": ""
    }
  }
]
```

#### 响应 (Response)

**200 OK - 成功**

| 字段 | 类型 | 描述 |
| :--- | :--- | :--- |
| `message` | string | 操作成功的提示信息。 |
| `filename` | string | 已上传文件的名称。 |
| `items_processed` | integer | JSON文件中被成功处理并入库的对象数量。 |

**示例:**
```json
{
  "message": "文件 'my_data.json' 上传并处理成功。",
  "filename": "my_data.json",
  "items_processed": 2
}
```

**错误响应**
* `400 Bad Request`: 如果文件不是 `.json` 格式、JSON内容无效、或JSON结构不符合规定（如顶层不是列表，或列表元素缺少`text`/`metadata`字段）。
* `503 Service Unavailable`: 如果后端核心模型未加载。
* `500 Internal Server Error`: 文件处理过程中发生其他错误。

#### cURL 示例
```bash
curl -X POST "[http://127.0.0.1:8000/upload_json](http://127.0.0.1:8000/upload_json)" \
-H "Content-Type: multipart/form-data" \
-F "file=@/path/to/your/my_data.json"
```
    """
    
    output_file = "信息检索_API_文档.docx"
    markdown_to_docx(markdown_content, output_file)
